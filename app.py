import os
import re
import qrcode
import random
import string
import tempfile
import base64
import streamlit as st
from datetime import datetime
from smtplib import SMTP
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email import encoders

# ✅ Aspose Words Cloud
from asposewordscloud import WordsApi
from asposewordscloud.models.requests import UploadFileRequest, SaveAsRequest, DownloadFileRequest
from asposewordscloud.models import PdfSaveOptionsData

# ✅ Google Sheets
from google.oauth2.service_account import Credentials
import gspread

# --- Streamlit Config ---
st.set_page_config("Completion Certificate Generator", layout="wide")

# --- Secrets ---
EMAIL = st.secrets["email"]["user"]
PASSWORD = st.secrets["email"]["password"]
APP_SID = st.secrets["aspose"]["app_sid"]
APP_KEY = st.secrets["aspose"]["app_key"]
SHEET_ID = st.secrets["gsheets"]["sheet_id"]
SHEET_NAME = st.secrets["gsheets"]["sheet_name"]

# --- Files ---
TEMPLATE_FILE = os.path.join(tempfile.gettempdir(), "completion_template.docx")
LOGO = "logo.png"

# --- Load Template from base64 ---
if not os.path.exists(TEMPLATE_FILE):
    encoded_template = st.secrets["template_base64"]["template_base64"]
    with open(TEMPLATE_FILE, "wb") as f:
        f.write(base64.b64decode(encoded_template))

# ✅ Aspose Setup
api = WordsApi(client_id=APP_SID, client_secret=APP_KEY)

# ✅ Google Sheets Setup
def get_gsheet():
    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive"
    ]
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=scopes
    )
    client = gspread.authorize(creds)
    return client.open_by_key(SHEET_ID).worksheet(SHEET_NAME)

def save_to_gsheet(data, status="Sent"):
    sheet = get_gsheet()
    row = [
        data['name'], data['domain'], data['month'],
        data['start_date'], data['end_date'],
        data['grade'], data['c_id'], data['email'], status
    ]
    sheet.append_row(row)

# ✅ Convert DOCX to PDF
def convert_to_pdf_asp(word_path, output_path):
    cloud_doc_name = os.path.basename(word_path)
    cloud_pdf_name = cloud_doc_name.replace(".docx", ".pdf")

    with open(word_path, "rb") as f:
        upload_result = api.upload_file(UploadFileRequest(f, cloud_doc_name))

    if not upload_result.uploaded or cloud_doc_name not in upload_result.uploaded:
        raise RuntimeError(f"Upload to Aspose failed. File {cloud_doc_name} not uploaded.")

    save_opts = PdfSaveOptionsData(file_name=cloud_pdf_name)
    api.save_as(SaveAsRequest(name=cloud_doc_name, save_options_data=save_opts))

    result = api.download_file(DownloadFileRequest(cloud_pdf_name))
    with open(output_path, "wb") as f:
        f.write(result)

# --- Utility Functions ---
def format_date(date_obj):
    return date_obj.strftime("%d %B %Y")

def generate_certificate_key():
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=9))

def generate_qr(data):
    qr = qrcode.QRCode(box_size=10, border=0)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    path = os.path.join(tempfile.gettempdir(), "qr.png")
    img.save(path)
    return path

def send_email(receiver, pdf_path, data):
    msg = MIMEMultipart()
    msg['From'] = EMAIL
    msg['To'] = receiver
    msg['Subject'] = f"🎉 Completion Certificate - {data['name']}"

    html = f"""
    <html><body>
        <p>Dear <strong>{data['name']}</strong>,</p>
        <p>Congratulations on completing your <strong>{data['month']} month</strong> internship at <strong>SkyHighes Technology</strong>!</p>
        <p><b>Details:</b></p>
        <ul>
            <li><strong>Domain:</strong> {data['domain']}</li>
            <li><strong>Duration:</strong> {data['start_date']} to {data['end_date']}</li>
            <li><strong>Grade:</strong> {data['grade']}</li>
            <li><strong>Certificate ID:</strong> {data['c_id']}</li>
        </ul>
        <p>Your certificate is attached as a PDF.</p>
        <p>All the best for your future!</p>
        <p><strong>SkyHighes Technology Team</strong></p>
    </body></html>
    """
    msg.attach(MIMEText(html, 'html'))

    with open(pdf_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={os.path.basename(pdf_path)}")
        msg.attach(part)

    with SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(EMAIL, PASSWORD)
        server.send_message(msg)

# --- Styling ---
st.markdown("""
<style>
.title-text {
    font-size: 2rem;
    font-weight: 700;
}
.stButton>button {
    background-color: #1E88E5;
    color: white;
    padding: 0.5rem 1.5rem;
    border-radius: 8px;
    font-weight: 600;
}
</style>
""", unsafe_allow_html=True)

# --- Header ---
with st.container():
    col_logo, col_title = st.columns([1, 6])
    with col_logo:
        if os.path.exists(LOGO):
            st.image(LOGO, width=80)
    with col_title:
        st.markdown('<div class="title-text">SkyHighes Technologies Completion Certificate Portal</div>', unsafe_allow_html=True)

st.divider()

# --- Form ---
with st.form("certificate_form"):
    st.subheader("🎓 Generate Completion Certificate")
    col1, col2, col3 = st.columns(3)
    with col1:
        name = st.text_input("Intern Name")
    with col2:
        domain = st.text_input("Domain")
    with col3:
        email = st.text_input("Recipient Email")

    col4, col5, col6 = st.columns(3)
    with col4:
        month = st.number_input("Internship Duration (Months)", min_value=1, max_value=12)
    with col5:
        start_date = st.date_input("Start Date", value=datetime.today())
    with col6:
        end_date = st.date_input("End Date", value=datetime.today())

    grade = st.selectbox("Grade", ["A+", "A", "B+", "B", "C"])
    submit = st.form_submit_button("🎯 Generate & Send Certificate")

# --- On Submit ---
if submit:
    if not all([name, domain, email]):
        st.error("❌ Please fill all fields.")
    elif end_date < start_date:
        st.warning("⚠️ End date cannot be before start date.")
    elif not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        st.warning("⚠️ Invalid email.")
    else:
        cert_id = generate_certificate_key()
        data = {
            "name": name.strip(),
            "domain": domain.strip(),
            "month": month,
            "start_date": format_date(start_date),
            "end_date": format_date(end_date),
            "grade": grade,
            "c_id": cert_id,
            "email": email.strip()
        }

        # Step 1: Insert QR into DOCX
        qr_path = generate_qr(f"{name}, {domain}, {month}, {data['start_date']}, {data['end_date']}, {grade}, {cert_id}")
        docx_raw = Document(TEMPLATE_FILE)
        try:
            cell = docx_raw.tables[0].rows[0].cells[0]
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.clear()
            run = para.add_run()
            run.add_picture(qr_path, width=Inches(1.42), height=Inches(1.42))
            qr_template = os.path.join(tempfile.gettempdir(), "template_with_qr.docx")
            docx_raw.save(qr_template)
        except Exception as e:
            st.warning(f"⚠️ QR insert failed: {e}")
            qr_template = TEMPLATE_FILE

        # Step 2: Render template
        doc = DocxTemplate(qr_template)
        doc.render(data)
        docx_path = os.path.join(tempfile.gettempdir(), f"Certificate_{name}.docx")
        doc.save(docx_path)

        # Step 3: Convert to PDF
        pdf_path = os.path.join(tempfile.gettempdir(), f"Certificate_{name}.pdf")
        try:
            convert_to_pdf_asp(docx_path, pdf_path)
        except Exception as e:
            st.error(f"❌ Aspose conversion failed: {e}")
            pdf_path = docx_path

        # Step 4a: Send Email
        try:
            send_email(email, pdf_path, data)
            st.success(f"✅ Certificate sent to {email}")
        except Exception as e:
            st.error(f"❌ Email sending failed: {e}")

        # ✅ Test GSheet connectivity
        try:
            sheet = get_gsheet()
            st.success(f"✅ Connected to Google Sheet: {sheet.title}")
        except Exception as e:
            st.error(f"❌ Google Sheet connection failed: {type(e).__name__}: {e}")
            st.write("👀 client_email:", st.secrets["gcp_service_account"]["client_email"])
            st.write("📄 sheet_id:", st.secrets["gsheets"]["sheet_id"])
            st.write("🗂️ sheet_name:", st.secrets["gsheets"]["sheet_name"])

        
        # Step 4b: Log to Google Sheet
        try:
            save_to_gsheet(data)
            st.success("✅ Logged to Google Sheet")
        except Exception as e:
            st.error(f"❌ Google Sheet logging failed: {type(e).__name__}: {e}")
        
        # Step 4c: Offer PDF Download
        try:
            with open(pdf_path, "rb") as f:
                st.download_button("📥 Download Certificate", f, file_name=os.path.basename(pdf_path))
        except Exception as e:
            st.warning(f"⚠️ PDF file could not be offered for download: {e}")


# --- Footer ---
st.markdown("<hr><center><small>© 2025 SkyHighes Technologies. All Rights Reserved.</small></center>", unsafe_allow_html=True)
